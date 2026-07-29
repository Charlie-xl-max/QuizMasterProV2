# -*- coding: utf-8 -*-
"""
题库解析模块 - 支持解析 docx、pdf、zip、txt 格式的题库文件
Copyright (c) 2026 QuizMasterPro V2 Contributors
Licensed under the MIT License (see LICENSE file for details)

支持的题库格式：
    判断题
    题目内容
    【答案】对/错

    单选题
    题目内容
    A. 选项A
    B. 选项B
    C. 选项C
    D. 选项D
    【答案】A

    多选题
    题目内容
    A. 选项A
    B. 选项B
    C. 选项C
    D. 选项D
    【答案】ABC

ZIP 支持按学科文件夹组织，系统自动识别学科名称
"""

import os
import re
import json
import zipfile
import tempfile
import shutil
from typing import Dict, List, Optional, Tuple, Any

# ── 可选依赖检测 ──────────────────────────────────────
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pdfplumber
    PDF_AVAILABLE = True
    PDF_BACKEND = 'pdfplumber'
except ImportError:
    try:
        import PyPDF2
        PDF_AVAILABLE = True
        PDF_BACKEND = 'PyPDF2'
    except ImportError:
        PDF_AVAILABLE = False
        PDF_BACKEND = None

# ── 常量 ──────────────────────────────────────────────
QUESTION_TYPE_MAP = {
    '判断题': 'judge',
    '单选题': 'single',
    '多选题': 'multi',
}

TRUE_ANSWERS = {'对', 'T', 't', '√', '正确', '是', 'True', 'true', '正确的', 'TRUE'}
FALSE_ANSWERS = {'错', 'F', 'f', '×', '错误', '否', 'False', 'false', 'x', 'X', '错误的', 'FALSE'}

SUPPORTED_EXTS = ('.docx', '.pdf', '.txt', '.zip')


def normalize_judge_answer(answer: str) -> str:
    """标准化判断题答案"""
    answer = answer.strip()
    if answer in TRUE_ANSWERS:
        return '对'
    if answer in FALSE_ANSWERS:
        return '错'
    return answer


# ── 文本提取 ──────────────────────────────────────────

def extract_text_from_docx(filepath: str) -> str:
    """从 Word (.docx) 文件提取文本"""
    if not DOCX_AVAILABLE:
        raise ImportError(
            "python-docx 未安装。请运行: pip install python-docx\n"
            "如使用 .exe 版本，此依赖已内置。"
        )
    doc = Document(filepath)
    lines = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text.strip())
    # 同时提取表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    lines.append(cell.text.strip())
    return '\n'.join(lines)


def extract_text_from_pdf(filepath: str) -> str:
    """从 PDF 文件提取文本"""
    if not PDF_AVAILABLE:
        raise ImportError(
            "pdfplumber 或 PyPDF2 未安装。请运行: pip install pdfplumber\n"
            "如使用 .exe 版本，此依赖已内置。"
        )

    text_parts = []

    if PDF_BACKEND == 'pdfplumber':
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
    elif PDF_BACKEND == 'PyPDF2':
        with open(filepath, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

    return '\n'.join(text_parts)


# ── 题目检测 ──────────────────────────────────────────

def detect_question_type(line: str) -> Optional[str]:
    """检测行是否为题型标记，返回题型 key"""
    line = line.strip()
    for type_name, type_key in QUESTION_TYPE_MAP.items():
        # 匹配: 判断题 / 【判断题】 / [判断题] / （判断题）等
        pattern = rf'^[【\[\（(]?\s*{type_name}\s*[】\]\）)]?\s*$'
        if re.match(pattern, line):
            return type_key
        # 也匹配 "一、判断题" 这种格式
        pattern2 = rf'^[一二三四五六七八九十\d]+[、，\.\s]*{type_name}\s*$'
        if re.match(pattern2, line):
            return type_key
    return None


def extract_question_type_name(line: str) -> Optional[str]:
    """提取题型的中文名"""
    line = line.strip()
    for type_name in QUESTION_TYPE_MAP:
        if type_name in line:
            return type_name
    return None


# ── 选项和答案解析 ────────────────────────────────────

def parse_options(text: str) -> Dict[str, str]:
    """从文本中解析选项 {A: '内容', B: '内容', ...}"""
    options = {}
    # 匹配 A. xxx B. xxx 或 A、xxx B、xxx
    pattern = r'([A-Z])[\.、\s]+(.+?)(?=[A-Z][\.、\s]|$)'
    matches = re.findall(pattern, text, re.DOTALL)
    for letter, content in matches:
        content = content.strip()
        if content:
            options[letter] = content
    return options


def extract_answer(line: str) -> Optional[str]:
    """从行中提取答案"""
    line = line.strip()
    patterns = [
        r'[【\[\（(]答案[】\]\）)]\s*[:：]?\s*(.+)',
        r'^答案\s*[:：]\s*(.+)',
        r'参考答案\s*[:：]\s*(.+)',
        r'正确答案\s*[:：]\s*(.+)',
    ]
    for pattern in patterns:
        match = re.search(pattern, line)
        if match:
            return match.group(1).strip()
    return None


def extract_stem_and_options(text: str) -> Tuple[str, Dict[str, str]]:
    """从文本中分离题干和选项"""
    option_pattern = r'[A-Z][\.、\s]'
    match = re.search(option_pattern, text)
    if not match:
        return text.strip(), {}

    stem = text[:match.start()].strip()
    options_text = text[match.start():].strip()
    options = parse_options(options_text)
    return stem, options


def clean_text(text: str) -> str:
    """清理文本：合并多余空白"""
    text = re.sub(r'\s+', ' ', text)
    text = text.strip()
    return text


# ── 核心解析逻辑 ──────────────────────────────────────

def parse_questions_from_text(text: str, subject_name: str) -> List[Dict[str, Any]]:
    """
    从文本中解析题目列表。
    格式要求：
        判断题
        题目内容
        【答案】对

        单选题
        题目内容
        A. 选项A
        B. 选项B
        【答案】A

        多选题
        题目内容
        A. 选项A
        B. 选项B
        C. 选项C
        D. 选项D
        【答案】ABC
    """
    lines = text.split('\n')
    questions = []
    current_q = None
    current_options_text = []
    collecting_options = False
    question_id = 0
    parse_errors = []

    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue

        q_type = detect_question_type(line)
        answer = extract_answer(line)

        if q_type:
            # 保存上一题
            if current_q is not None:
                _finalize_question(current_q, current_options_text, collecting_options)
                if _is_valid_question(current_q):
                    questions.append(current_q)
                else:
                    parse_errors.append(
                        f"第{len(questions)+1}题不完整：{current_q.get('stem','')[:30]}..."
                    )

            question_id += 1
            type_name = extract_question_type_name(line) or ''
            stem_text = re.sub(
                rf'^[【\[\（(]?\s*{re.escape(type_name)}\s*[】\]\）)]?\s*', '', line
            )
            stem_text = re.sub(r'^[一二三四五六七八九十\d]+[、，\.\s]+', '', stem_text)

            current_q = {
                'id': question_id,
                'stem': '',
                'options': {},
                'answer': '',
                'type': q_type,
                'subject': subject_name,
            }
            current_options_text = []
            collecting_options = False

            if q_type == 'judge':
                current_q['stem'] = clean_text(stem_text) if stem_text else ''
            else:
                stem, options = extract_stem_and_options(stem_text)
                current_q['stem'] = clean_text(stem)
                current_q['options'] = options
                if options:
                    collecting_options = True

        elif answer is not None and current_q is not None:
            raw_answer = answer
            if current_q['type'] == 'judge':
                current_q['answer'] = normalize_judge_answer(raw_answer)
            else:
                raw_answer = re.sub(r'[\s,，、]+', '', raw_answer).upper()
                current_q['answer'] = raw_answer

            if collecting_options and current_options_text:
                options_full = ' '.join(current_options_text)
                existing_options = current_q.get('options', {})
                new_options = parse_options(options_full)
                existing_options.update(new_options)
                current_q['options'] = existing_options

            collecting_options = False

        elif current_q is not None:
            if current_q['type'] == 'judge':
                if not current_q['stem']:
                    current_q['stem'] = line
                else:
                    current_q['stem'] += ' ' + line
            else:
                has_option_marker = bool(re.match(r'^[A-Z][\.、\s]', line))
                if has_option_marker or collecting_options:
                    collecting_options = True
                    current_options_text.append(line)
                else:
                    if not current_q['stem']:
                        current_q['stem'] = line
                    else:
                        current_q['stem'] += ' ' + line

    # 处理最后一题
    if current_q is not None:
        _finalize_question(current_q, current_options_text, collecting_options)
        if _is_valid_question(current_q):
            questions.append(current_q)
        else:
            parse_errors.append(
                f"第{len(questions)+1}题不完整：{current_q.get('stem','')[:30]}..."
            )

    # 打印解析警告
    if parse_errors:
        print(f"解析警告 ({subject_name}):")
        for err in parse_errors[:5]:
            print(f"  - {err}")
        if len(parse_errors) > 5:
            print(f"  ... 还有 {len(parse_errors)-5} 个警告")

    return questions


def _finalize_question(q: dict, options_text: list, collecting: bool):
    """完成题目解析的收尾工作"""
    if collecting and options_text:
        options_full = ' '.join(options_text)
        existing_options = q.get('options', {})
        new_options = parse_options(options_full)
        existing_options.update(new_options)
        q['options'] = existing_options


def _is_valid_question(q: dict) -> bool:
    """检查题目是否有效"""
    if not q.get('stem'):
        return False
    if not q.get('answer'):
        return False
    if q['type'] in ('single', 'multi') and not q.get('options'):
        return False
    return True


# ── 文件名处理 ────────────────────────────────────────

def get_subject_name_from_filename(filename: str) -> str:
    """从文件名提取学科名称"""
    name = os.path.splitext(os.path.basename(filename))[0]
    # 去掉常见后缀
    for suffix in ['-判断题', '-单选题', '-多选题', '_判断题', '_单选题', '_多选题',
                   '判断题', '单选题', '多选题', '-题目', '_题目', '-题库', '_题库']:
        if name.endswith(suffix):
            name = name[:-len(suffix)]
            break
    return name.strip('_- ')


def generate_subject_id(subject_name: str) -> str:
    """为学科名称生成英文 ID"""
    pinyin_map = {
        '计算机网络': 'network',
        '数据结构': 'datastruct',
        '操作系统': 'os',
        '数据库原理': 'database',
        '软件工程': 'swe',
        '编译原理': 'compiler',
        '计算机组成原理': 'comporg',
        '线性代数': 'linalg',
        '概率论': 'probability',
        '英语': 'english',
    }

    if subject_name in pinyin_map:
        return pinyin_map[subject_name]

    # 尝试提取英文/拼音
    simple = re.sub(r'[^\w]', '', subject_name.lower())
    if simple:
        return simple[:20]

    return f"subject_{hash(subject_name) % 10000}"


# ── 单文件解析 ────────────────────────────────────────

def parse_file(filepath: str, subject_id: str = None) -> Dict[str, Any]:
    """解析单个文件"""
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"文件不存在: {filepath}")

    ext = os.path.splitext(filepath)[1].lower()
    if ext not in SUPPORTED_EXTS:
        raise ValueError(
            f"不支持的文件格式: {ext}。"
            f"支持的格式: {', '.join(SUPPORTED_EXTS)}"
        )

    subj_name = get_subject_name_from_filename(filepath)
    if not subject_id:
        subject_id = generate_subject_id(subj_name)

    try:
        if ext == '.docx':
            text = extract_text_from_docx(filepath)
        elif ext == '.pdf':
            text = extract_text_from_pdf(filepath)
        elif ext == '.txt':
            with open(filepath, 'r', encoding='utf-8') as f:
                text = f.read()
        else:
            raise ValueError(f"不支持的文件格式: {ext}")

        if not text.strip():
            raise ValueError(
                f"文件内容为空，请检查文件是否正确。\n"
                f"提示：扫描版 PDF 无法直接提取文字，请使用文字版 PDF。"
            )

        questions = parse_questions_from_text(text, subject_id)
        return {
            'subject_id': subject_id,
            'subject_name': subj_name,
            'questions': questions,
        }
    except (ImportError, ValueError) as e:
        raise
    except Exception as e:
        raise RuntimeError(
            f"解析文件失败 {os.path.basename(filepath)}:\n{str(e)}"
        )


# ── ZIP 解析 ──────────────────────────────────────────

def parse_zip(zip_path: str) -> List[Dict[str, Any]]:
    """解析 ZIP 压缩包"""
    if not os.path.exists(zip_path):
        raise FileNotFoundError(f"文件不存在: {zip_path}")

    results = []
    temp_dir = tempfile.mkdtemp()

    try:
        with zipfile.ZipFile(zip_path, 'r') as zf:
            zf.extractall(temp_dir)

        # 检测是否有子文件夹（按学科组织）
        subject_dirs = []
        for item in os.listdir(temp_dir):
            item_path = os.path.join(temp_dir, item)
            if os.path.isdir(item_path) and not item.startswith('__'):
                subject_dirs.append(item)

        if subject_dirs:
            # 按学科文件夹组织
            for subj_dir in sorted(subject_dirs):
                subj_path = os.path.join(temp_dir, subj_dir)
                all_questions = []

                for root, dirs, files in os.walk(subj_path):
                    for fname in sorted(files):
                        fpath = os.path.join(root, fname)
                        ext = os.path.splitext(fname)[1].lower()
                        if ext in ('.docx', '.pdf', '.txt'):
                            try:
                                result = parse_file(fpath)
                                all_questions.extend(result['questions'])
                            except Exception as e:
                                print(f"解析压缩包内文件失败 {fname}: {e}")

                if all_questions:
                    for q in all_questions:
                        q['subject'] = subj_dir
                    results.append({
                        'subject_name': subj_dir,
                        'questions': all_questions,
                    })
        else:
            # 扁平结构：所有文件合并为一个科目
            zip_name = os.path.splitext(os.path.basename(zip_path))[0]
            all_questions = []
            for root, dirs, files in os.walk(temp_dir):
                for fname in sorted(files):
                    fpath = os.path.join(root, fname)
                    ext = os.path.splitext(fname)[1].lower()
                    if ext in ('.docx', '.pdf', '.txt'):
                        try:
                            result = parse_file(fpath)
                            all_questions.extend(result['questions'])
                        except Exception as e:
                            print(f"解析压缩包内文件失败 {fname}: {e}")

            if all_questions:
                results.append({
                    'subject_name': zip_name,
                    'questions': all_questions,
                })

    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)

    return results


# ── 文件夹扫描 ────────────────────────────────────────

def _find_subject_dirs(root_dir: str) -> List[str]:
    """递归查找包含题库文件的学科文件夹"""
    subj_dirs = []
    for item in sorted(os.listdir(root_dir)):
        item_path = os.path.join(root_dir, item)
        if os.path.isdir(item_path) and not item.startswith('__'):
            sub_subj = _find_subject_dirs(item_path)
            if sub_subj:
                subj_dirs.extend(sub_subj)
            else:
                has_q_files = False
                try:
                    for f in os.listdir(item_path):
                        if f.lower().endswith(('.docx', '.pdf', '.txt')):
                            has_q_files = True
                            break
                except Exception:
                    pass
                if has_q_files:
                    subj_dirs.append(item_path)
    return subj_dirs


def parse_folder(folder_path: str) -> Dict[str, Any]:
    """解析文件夹中的所有题库文件"""
    if not os.path.exists(folder_path):
        raise FileNotFoundError(f"文件夹不存在: {folder_path}")

    subjects_data = []
    seen_ids = set()

    def _add_subject(subj_id, subj_name, questions):
        if not questions:
            return
        # 去重 ID
        if subj_id in seen_ids:
            orig = subj_id
            c = 2
            while f"{orig}_{c}" in seen_ids:
                c += 1
            subj_id = f"{orig}_{c}"
        seen_ids.add(subj_id)

        # 按题型分组并分配 ID
        by_type = {'judge': [], 'single': [], 'multi': []}
        for q in questions:
            q['subject'] = subj_id
            if q['type'] in by_type:
                by_type[q['type']].append(q)
            else:
                by_type['single'].append(q)

        final_qs = []
        for qtype in ['judge', 'single', 'multi']:
            for i, q in enumerate(by_type[qtype]):
                q['id'] = f"{subj_id}_{qtype}_{i+1:04d}"
                final_qs.append(q)

        subjects_data.append({
            'id': subj_id,
            'name': subj_name,
            'questions': final_qs,
        })

    # 查找学科子文件夹
    subj_dirs = _find_subject_dirs(folder_path)

    if subj_dirs:
        for subj_dir in subj_dirs:
            subj_name = os.path.basename(subj_dir)
            subj_id = generate_subject_id(subj_name)
            subj_questions = []
            for root, dirs, files in os.walk(subj_dir):
                for fname in sorted(files):
                    fpath = os.path.join(root, fname)
                    ext = os.path.splitext(fname)[1].lower()
                    if ext in ('.docx', '.pdf', '.txt'):
                        try:
                            result = parse_file(fpath, subj_id)
                            subj_questions.extend(result['questions'])
                        except Exception as e:
                            print(f"解析文件失败 {fname}: {e}")
                    elif ext == '.zip':
                        try:
                            zip_results = parse_zip(fpath)
                            for zr in zip_results:
                                subj_questions.extend(zr['questions'])
                        except Exception as e:
                            print(f"解析压缩包失败 {fname}: {e}")
            _add_subject(subj_id, subj_name, subj_questions)
    else:
        # 扁平结构：直接扫描文件
        for item in sorted(os.listdir(folder_path)):
            item_path = os.path.join(folder_path, item)
            if not os.path.isfile(item_path):
                continue
            ext = os.path.splitext(item)[1].lower()
            if ext not in SUPPORTED_EXTS:
                continue
            try:
                if ext == '.zip':
                    zip_results = parse_zip(item_path)
                    for zr in zip_results:
                        _add_subject(
                            generate_subject_id(zr['subject_name']),
                            zr['subject_name'],
                            zr['questions']
                        )
                else:
                    result = parse_file(item_path)
                    _add_subject(
                        generate_subject_id(result['subject_name']),
                        result['subject_name'],
                        result['questions']
                    )
            except Exception as e:
                print(f"处理文件失败 {item}: {e}")

    return {'subjects': subjects_data}


# ── 输出 ──────────────────────────────────────────────

def generate_questions_js(data: Dict[str, Any], output_path: str) -> None:
    """生成 questions.js 文件"""
    json_str = json.dumps(data, ensure_ascii=False)
    js_content = (
        "// 题库数据 - 由 QuizMasterPro V2 自动生成\n"
        "// 请通过「上传题库」功能更新，不要手动编辑此文件\n"
        "window.QUESTION_BANK = " + json_str + ";\n"
    )

    out_dir = os.path.dirname(os.path.abspath(output_path))
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(js_content)


# ── CLI ───────────────────────────────────────────────

if __name__ == '__main__':
    import sys

    if len(sys.argv) > 1:
        target = sys.argv[1]
    else:
        print("用法: python question_parser.py <文件或文件夹路径>")
        print("支持格式: .docx, .pdf, .txt, .zip")
        sys.exit(1)

    if os.path.isdir(target):
        result = parse_folder(target)
        output_file = os.path.join(target, 'questions.js')
        generate_questions_js(result, output_file)
        print(f"解析完成！共 {len(result['subjects'])} 个学科")
        for subj in result['subjects']:
            print(f"  - {subj['name']}: {len(subj['questions'])} 题")
        print(f"已输出到: {output_file}")
    elif os.path.isfile(target):
        ext = os.path.splitext(target)[1].lower()
        if ext == '.zip':
            results = parse_zip(target)
            print(f"解析完成！共 {len(results)} 个学科")
            for r in results:
                print(f"  - {r['subject_name']}: {len(r['questions'])} 题")
        else:
            result = parse_file(target)
            print(f"解析完成！学科: {result['subject_name']}, 共 {len(result['questions'])} 题")
    else:
        print(f"路径不存在: {target}")
        sys.exit(1)
